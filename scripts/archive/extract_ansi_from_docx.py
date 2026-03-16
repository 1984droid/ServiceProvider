"""
Extract Images and Tables from ANSI A92.2-2021 DOCX

This script extracts:
1. All 17 embedded images (figures, logos, symbols)
2. All 17 tables as formatted PNG images

Source: ANSI+SAIA+A92.2-2021.docx
Output: static/inspection_references/ansi_a92_2_2021/
"""

import os
import io
from pathlib import Path
from docx import Document
from PIL import Image, ImageDraw, ImageFont

# Image mapping based on media file names in DOCX
# Format: 'media/imageX.ext': ('subdir', 'output_filename')
IMAGE_MAPPING = {
    'media/image1.png': ('logos', 'ansi_logo.png'),
    'media/image2.jpeg': ('figures', 'figure_1_dielectric_test_cat_ab.jpg'),
    'media/image3.jpeg': ('figures', 'figure_1a_electrode_assembly_details.jpg'),
    'media/image4.png': ('figures', 'figure_2_dielectric_test_cat_cde.png'),
    'media/image5.jpeg': ('figures', 'figure_2a_optional_test_config.jpg'),
    'media/image6.png': ('figures', 'figure_3_chassis_insulating_systems.png'),
    'media/image7.jpeg': ('figures', 'figure_3a_shunting_arrangement.jpg'),
    'media/image8.jpeg': ('figures', 'figure_4_boom_positions_test.jpg'),
    'media/image9.jpeg': ('figures', 'figure_5_bonding_arrangements.jpg'),
    'media/image10.png': ('figures', 'figure_6_upper_control_test.png'),
    'media/image11.jpeg': ('symbols', 'equipment_symbols_1.jpg'),
    'media/image12.jpeg': ('symbols', 'equipment_symbols_2.jpg'),
    'media/image13.jpeg': ('symbols', 'equipment_symbols_3.jpg'),
    'media/image14.jpeg': ('symbols', 'equipment_symbols_4.jpg'),
    'media/image15.jpeg': ('symbols', 'equipment_symbols_5.jpg'),
    'media/image16.jpeg': ('symbols', 'equipment_symbols_6.jpg'),
    'media/image17.jpeg': ('logos', 'saia_logo.jpg'),
}

# Table titles for identification
TABLE_TITLES = {
    1: 'Title Page',
    2: 'TABLE 2 - Periodic Electrical Test Values (Cat A&B with Lower Electrode)',
    3: 'TABLE 2 - Periodic Electrical Test Values (Cat C,D,E without Lower Electrode)',
    4: 'TABLE 2 - Periodic Electrical Test Values (Insulating Ladders and Towers)',
    5: 'TABLE 2 - Periodic Electrical Test Values (Insulating Ladders)',
    6: 'TABLE 1 - Design and Qualification Test Values (Cat A&B)',
    7: 'TABLE 1 - Design and Qualification Test Values (Cat C,D,E)',
    8: 'TABLE 1 - Design and Qualification Test Values (Ladders and Towers)',
    9: 'TABLE 4 - Test Voltages for Insulating Liners',
    10: 'TABLE 5 - Test Voltages for Insulating Buckets',
    11: 'Supplemental Table 1',
    12: 'Supplemental Table 2',
    13: 'TABLE - Equipment Categories and Applications',
    14: 'TABLE - Test Requirements by Category',
    15: 'TABLE 3 - Overload Test Line Contact Values',
    16: 'TABLE - Minimum Approach Distances',
    17: 'TABLE - Load Test Documentation',
}


def extract_images(docx_path: str, output_dir: str):
    """Extract all embedded images from DOCX."""
    print("Extracting images from DOCX...")

    doc = Document(docx_path)
    base_dir = Path(output_dir)

    images_extracted = []

    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            # Get image data
            image_data = rel.target_part.blob
            image_ref = rel.target_ref

            # Get mapping by target_ref
            if image_ref in IMAGE_MAPPING:
                subdir, filename = IMAGE_MAPPING[image_ref]
                output_path = base_dir / subdir / filename

                # Save image
                output_path.parent.mkdir(parents=True, exist_ok=True)
                with open(output_path, 'wb') as f:
                    f.write(image_data)

                images_extracted.append((image_ref, subdir, filename))
                print(f"  {subdir:8s} -> {filename}")

    return images_extracted


def create_table_image(table, title: str, output_path: Path):
    """Convert DOCX table to formatted PNG image."""
    # Extract table data
    table_data = []
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        table_data.append(row_data)

    if not table_data:
        return False

    # Calculate image dimensions
    font_size = 14
    cell_padding = 10
    line_height = font_size + cell_padding
    title_height = 40

    # Calculate column widths based on content
    num_cols = len(table_data[0])
    col_widths = [0] * num_cols

    for row in table_data:
        for i, cell in enumerate(row):
            if i < len(col_widths):
                # Estimate width (rough approximation)
                cell_width = len(cell) * (font_size // 2) + cell_padding * 2
                col_widths[i] = max(col_widths[i], cell_width)

    # Set maximum column width
    max_col_width = 300
    col_widths = [min(w, max_col_width) for w in col_widths]

    table_width = sum(col_widths) + (num_cols + 1) * 2  # 2px borders
    table_height = len(table_data) * line_height + title_height + (len(table_data) + 1) * 2

    # Create image
    img = Image.new('RGB', (table_width, table_height), color='white')
    draw = ImageDraw.Draw(img)

    # Try to use a better font, fall back to default
    try:
        title_font = ImageFont.truetype("arial.ttf", 16)
        cell_font = ImageFont.truetype("arial.ttf", font_size)
    except:
        title_font = ImageFont.load_default()
        cell_font = ImageFont.load_default()

    # Draw title
    draw.text((10, 10), title, fill='black', font=title_font)

    # Draw table
    y_offset = title_height

    for row_idx, row in enumerate(table_data):
        x_offset = 0

        for col_idx, cell in enumerate(row):
            if col_idx >= len(col_widths):
                break

            # Draw cell border
            x1 = x_offset
            y1 = y_offset
            x2 = x_offset + col_widths[col_idx]
            y2 = y_offset + line_height

            # Header row background
            if row_idx == 0:
                draw.rectangle([x1, y1, x2, y2], fill='#E0E0E0', outline='black')
            else:
                draw.rectangle([x1, y1, x2, y2], outline='black')

            # Draw text (truncate if too long)
            cell_text = cell[:50] if len(cell) > 50 else cell
            draw.text((x1 + 5, y1 + 5), cell_text, fill='black', font=cell_font)

            x_offset += col_widths[col_idx]

        y_offset += line_height

    # Save image
    output_path.parent.mkdir(parents=True, exist_ok=True)
    img.save(output_path, 'PNG', optimize=True)

    return True


def extract_tables(docx_path: str, output_dir: str):
    """Extract all tables as PNG images."""
    print("\nExtracting tables as PNG images...")

    doc = Document(docx_path)
    base_dir = Path(output_dir) / 'tables'

    tables_extracted = []

    for idx, table in enumerate(doc.tables, start=1):
        title = TABLE_TITLES.get(idx, f'Table {idx}')
        filename = f"table_{idx:02d}.png"
        output_path = base_dir / filename

        success = create_table_image(table, title, output_path)

        if success:
            tables_extracted.append((idx, title, filename))
            print(f"  [{idx:2d}] {filename} - {title}")

    return tables_extracted


def create_readme(output_dir: str, images: list, tables: list):
    """Create comprehensive README documenting all extracted resources."""

    readme_content = """# ANSI A92.2-2021 Reference Materials

**Source:** ANSI+SAIA+A92.2-2021.docx
**Standard:** American National Standard for Vehicle-Mounted Elevating and Rotating Aerial Devices
**Extracted:** Figures, Tables, Logos, and Equipment Symbols

---

## Directory Structure

```
ansi_a92_2_2021/
├── figures/     (9 images)  - Technical diagrams from standard
├── tables/      (17 images) - All tables as formatted PNG images
├── logos/       (2 images)  - ANSI and SAIA logos
├── symbols/     (6 images)  - Equipment warning and operational symbols
└── README.md
```

---

## Figures (9)

Technical diagrams referenced throughout the ANSI A92.2-2021 standard.

### Dielectric Testing
- `figure_1_dielectric_test_cat_ab.jpg` - **Figure 1**: Test setup for Category A/B devices with lower electrode
- `figure_1a_electrode_assembly_details.jpg` - **Figure 1A**: Lower test electrode assembly details
- `figure_2_dielectric_test_cat_cde.png` - **Figure 2**: Test setup for Category C/D/E devices
- `figure_2a_optional_test_config.jpg` - **Figure 2A**: Optional test configuration

### Chassis and Bonding
- `figure_3_chassis_insulating_systems.png` - **Figure 3**: Chassis and insulating systems
- `figure_3a_shunting_arrangement.jpg` - **Figure 3A**: Shunting arrangement for energized line contact
- `figure_5_bonding_arrangements.jpg` - **Figure 5**: Electrical bonding arrangements

### Testing Procedures
- `figure_4_boom_positions_test.jpg` - **Figure 4**: Boom positions during load testing
- `figure_6_upper_control_test.png` - **Figure 6**: Upper control high resistance test setup

---

## Tables (17)

All tables extracted as formatted PNG images for easy reference.

### Design and Qualification Tests
- `table_06.png` - **TABLE 1**: Design/Qualification Test Values (Category A & B with lower electrode)
- `table_07.png` - **TABLE 1**: Design/Qualification Test Values (Category C, D, E without lower electrode)
- `table_08.png` - **TABLE 1**: Design/Qualification Test Values (Insulating Ladders and Towers)

### Periodic Electrical Tests
- `table_02.png` - **TABLE 2**: Periodic Test Values (Category A & B with lower electrode)
- `table_03.png` - **TABLE 2**: Periodic Test Values (Category C, D, E without lower electrode)
- `table_04.png` - **TABLE 2**: Periodic Test Values (Insulating Ladders and Towers)
- `table_05.png` - **TABLE 2**: Periodic Test Values (Insulating Ladders)

### Overload and Line Contact Tests
- `table_15.png` - **TABLE 3**: Overload Test Line Contact Values

### Component-Specific Tests
- `table_09.png` - **TABLE 4**: Test Voltages for Insulating Liners
- `table_10.png` - **TABLE 5**: Test Voltages for Insulating Buckets

### Equipment Categories and Applications
- `table_13.png` - Equipment Categories (A, B, C, D, E, Non-Insulating)
- `table_14.png` - Test Requirements by Category
- `table_16.png` - Minimum Approach Distances

### Supplemental Tables
- `table_01.png` - Title Page
- `table_11.png` - Supplemental Table 1
- `table_12.png` - Supplemental Table 2
- `table_17.png` - Load Test Documentation

---

## Logos (2)

- `ansi_logo.png` - American National Standards Institute logo
- `saia_logo.jpg` - Scaffold & Access Industry Association logo

**Usage:** Include in PDF reports and inspection documentation for professional appearance.

---

## Equipment Symbols (6)

Warning signs and operational pictograms that appear on aerial devices:

- `equipment_symbols_1.jpg` - High voltage warning signs
- `equipment_symbols_2.jpg` - Boom movement indicators
- `equipment_symbols_3.jpg` - Operational control symbols
- `equipment_symbols_4.jpg` - Safety pictograms
- `equipment_symbols_5.jpg` - Capacity rating markers
- `equipment_symbols_6.jpg` - Operational marking symbols

**Usage:** Reference during markings and placards inspection steps.

---

## Usage in Inspection Templates

### Reference Images in Templates

```json
{
  "step_key": "dielectric_test_setup",
  "step_number": 1,
  "title": "Dielectric Test Setup - Category A/B",
  "reference_images": [
    {
      "image_id": "figure_1",
      "title": "Figure 1: Test Setup for Category A/B",
      "file_path": "ansi_a92_2_2021/figures/figure_1_dielectric_test_cat_ab.jpg",
      "caption": "Lower test electrode system configuration per Section 5.4.2",
      "display_mode": "inline",
      "figure_number": "Figure 1"
    },
    {
      "image_id": "table_2_cat_ab",
      "title": "TABLE 2: Periodic Test Values",
      "file_path": "ansi_a92_2_2021/tables/table_02.png",
      "caption": "Test voltage and current limits for Category A/B devices",
      "display_mode": "modal",
      "figure_number": "Table 2"
    }
  ]
}
```

### Display Modes
- **inline**: Image always visible on inspection screen
- **modal**: Click/tap to view full size
- **both**: Thumbnail inline + click for full size

---

## File Formats

- **Figures**: JPEG/PNG (original format from standard)
- **Tables**: PNG (generated from DOCX tables)
- **Logos**: PNG/JPEG (original format)
- **Symbols**: JPEG (original format)

---

## Notes

- All images extracted from official ANSI+SAIA+A92.2-2021.docx document
- Tables are rendered as images for consistent display across platforms
- Figure numbers correspond to ANSI A92.2-2021 standard references
- Images optimized for web/mobile display while maintaining readability

---

## Related Files

- Source: `ANSI+SAIA+A92.2-2021.docx`
- Extraction Script: `scripts/extract_ansi_from_docx.py`
- Template Schema: `apps/inspections/schemas/template_schema.py`
"""

    readme_path = Path(output_dir) / 'README.md'
    with open(readme_path, 'w', encoding='utf-8') as f:
        f.write(readme_content)

    print(f"\nREADME created: {readme_path}")


def main():
    """Main extraction process."""
    docx_path = 'ANSI+SAIA+A92.2-2021.docx'
    output_dir = 'static/inspection_references/ansi_a92_2_2021'

    print("="*80)
    print("ANSI A92.2-2021 DOCX Extraction")
    print("="*80)
    print(f"\nSource: {docx_path}")
    print(f"Output: {output_dir}\n")

    # Extract images
    images = extract_images(docx_path, output_dir)
    print(f"\n-> Extracted {len(images)} images")

    # Extract tables
    tables = extract_tables(docx_path, output_dir)
    print(f"\n-> Extracted {len(tables)} tables")

    # Create README
    create_readme(output_dir, images, tables)

    print("\n" + "="*80)
    print("EXTRACTION COMPLETE")
    print("="*80)
    print(f"\nFigures:  {len([i for i in images if i[1] == 'figures'])}")
    print(f"Logos:    {len([i for i in images if i[1] == 'logos'])}")
    print(f"Symbols:  {len([i for i in images if i[1] == 'symbols'])}")
    print(f"Tables:   {len(tables)}")
    print(f"\nAll resources ready for inspection template integration!")


if __name__ == '__main__':
    main()
